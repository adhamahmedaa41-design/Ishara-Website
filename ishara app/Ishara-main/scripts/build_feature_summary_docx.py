"""Generate Ishara_Feature_Summary.docx from the in-conversation summary.
Run: python scripts/build_feature_summary_docx.py
Output: ./Ishara_Feature_Summary.docx
"""

import os, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.environ["PYTHONIOENCODING"] = "utf-8"
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Ishara_Feature_Summary.docx"

TEAL = RGBColor(0x0E, 0x7C, 0x86)
ORANGE = RGBColor(0xE0, 0x6B, 0x2A)
MUTED = RGBColor(0x55, 0x55, 0x55)
INK = RGBColor(0x1A, 0x1A, 0x1A)


def style_normal(doc: Document):
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = INK
    s.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = TEAL


def add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(14)


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = TEAL


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = ORANGE


def add_para(doc: Document, runs):
    """runs: list of (text, {bold, italic, code}) tuples."""
    p = doc.add_paragraph()
    for text, opts in runs:
        r = p.add_run(text)
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", False)
        if opts.get("code"):
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            r.font.size = Pt(11)
    return p


def add_text(doc: Document, text: str):
    add_para(doc, [(text, {})])


def add_bullet(doc: Document, runs, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    for text, opts in runs:
        r = p.add_run(text)
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", False)
        if opts.get("code"):
            r.font.name = "Consolas"
            r.font.size = Pt(10)


def add_table(doc: Document, header, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def code_inline(text):
    return (text, {"code": True})


def b(text):
    return (text, {"bold": True})


def n(text):
    return (text, {})


def i(text):
    return (text, {"italic": True})


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style_normal(doc)

    # ── Title ────────────────────────────────────────────────────────────
    add_title(doc, "Ishara")
    add_subtitle(doc, "Accessible ESL assistant and safety companion — full feature inventory")

    # ── Intro ────────────────────────────────────────────────────────────
    add_para(doc, [n(
        "Below is every shipping feature in the app, grouped by area, with the screen it lives on, "
        "the on-device pipeline it uses, and what the user actually sees and does."
    )])

    # ── §1 Auth & onboarding ─────────────────────────────────────────────
    add_h1(doc, "1. Auth & onboarding flow")
    add_table(doc,
        ["Screen", "Path", "Purpose"],
        [
            ["Splash", "/", "Boot-time auth check. If authProvider.canUseApp → /home, else → /login."],
            ["Login", "/login", "Email + password against MongoDB Atlas via dio. Google + Facebook social login wired. Guest mode available."],
            ["Register", "/register", "New account creation → triggers email OTP."],
            ["OTP verify", "/otp", "Email verification step. On success routes to Profile-photo setup."],
            ["Profile-photo setup", "/profile-photo-setup", "One-time skippable step after OTP. Choose photo / take photo / skip. Default avatar from ui-avatars.com."],
            ["Server config", "/server-config", "Lets a user point the app at a different backend (LAN host / staging)."],
        ],
        col_widths=[1.4, 1.7, 3.3],
    )
    add_para(doc, [
        n("Credentials are stored in "), code_inline("flutter_secure_storage"),
        n("; non-secret session metadata in "), code_inline("shared_preferences"),
        n(". Auth state lives in a Riverpod "), code_inline("AuthState"),
        n(" that GoRouter's "), code_inline("refreshListenable"),
        n(" watches, so logout instantly bounces you back to /login."),
    ])

    # ── §2 Communicate ──────────────────────────────────────────────────
    add_h1(doc, "2. Communicate — bidirectional sign translator")
    add_para(doc, [
        n("Lives at "), code_inline("/home"),
        n(" (the Communicate screen is the default landing tab). Two directions, switched via a pill toggle."),
    ])

    add_h2(doc, "ESL → Arabic (camera mode)")
    add_para(doc, [n("The deaf user signs in front of the phone, the app emits Arabic words.")])
    add_bullet(doc, [n("The in-screen camera card streams frames into "), code_inline("sign_language_service.dart"),
                     n(", which uses ML Kit Pose detection + the on-device classifier model (loaded from "),
                     code_inline("assets/models/manifest.json"),
                     n(") with a 30-frame rolling buffer and a temporal smoother.")])
    add_bullet(doc, [n("Top-1 prediction is gated by a confidence floor and a per-word cooldown so the same sign doesn't fire 10× in a row.")])
    add_bullet(doc, [n("Live UI chips show: model-ready state, current FPS (target ≥ 6), frames collected vs. needed (N/30), low-light warning. A linear progress bar fills as the buffer fills, then flashes green when a detection lands.")])
    add_bullet(doc, [n("A "), b("\"Live Sign Translator\""), n(" FAB jumps to a dedicated full-screen YOLO live screen ("),
                     code_inline("/translator/yolo-live"),
                     n(") backed by "), code_inline("sign_yolo_detector.dart"),
                     n(" and "), code_inline("assets/models/sign_yolo.tflite"),
                     n(" (Arabic-10 classes, mAP50 0.995). Includes a dev-log overlay and lower confidence threshold for harder-to-catch signs.")])
    add_bullet(doc, [n("A "), b("\"Reset model\""), n(" action disposes and recreates the TFLite interpreter without restarting the camera — escape hatch when the on-device session wedges.")])
    add_bullet(doc, [n("A circuit-breaker recovery in the service handles Samsung-style out-of-order camera frames: first failure drops 6 frames, second cools 500 ms, third does a full image-stream restart. Resets after 30 clean frames.")])

    add_h2(doc, "Arabic → ESL (text + mic mode)")
    add_para(doc, [n("The hearing user types or speaks Arabic, the app shows the sign clips.")])
    add_bullet(doc, [n("Free-text input or "), code_inline("speech_to_text"), n(" voice input (mic button → live transcription).")])
    add_bullet(doc, [code_inline("psl_video_service.dart"),
                     n(" resolves each Arabic word against a curated PSL dictionary ("),
                     code_inline("assets/psl/pk-dictionary-mapping.json"),
                     n(" + "), code_inline("pk-dictionary-urls.json"),
                     n(") and streams the matching .mp4 from GitHub Releases via "),
                     code_inline("video_player"), n(".")])
    add_bullet(doc, [n("Each word renders as its own looping muted clip; words without a clip show a \"No sign clip for this word\" placeholder.")])
    add_bullet(doc, [n("A \"Speak\" button reads the translated text aloud via "), code_inline("flutter_tts"), n(".")])

    add_h2(doc, "Text → Sign")
    add_para(doc, [n("Standalone screen at "), code_inline("/translator/text-to-sign"),
                   n(" — same translation without the mic, useful in noisy contexts.")])

    # ── §3 Vision ───────────────────────────────────────────────────────
    add_h1(doc, "3. Vision — multi-mode camera understanding")
    add_para(doc, [n("Lives at "), code_inline("/vision"),
                   n(". Mode-switched via header chips; each mode is a self-contained pipeline.")])

    add_h2(doc, "Currency mode (EGP)")
    add_para(doc, [n("Point the camera at a banknote, the app says the total in Arabic.")])
    add_bullet(doc, [code_inline("assets/models/currency_egp.tflite"),
                     n(" — 6-class YOLOv8n (5 / 10 / 20 / 50 / 100 / 200 EGP, mAP50 0.995, md5 "),
                     code_inline("453eb1341a9abc77a16707609115a775"), n(").")])
    add_bullet(doc, [code_inline("currency_yolo_detector.dart"), n(" runs five stages:")])
    add_bullet(doc, [n("1. YOLO inference with input-shape autodetect (both NMS-true [1,N,6] and anchor-transposed [1,N,4+nc]).")], level=1)
    add_bullet(doc, [n("2. IoU NMS with class-aware suppression.")], level=1)
    add_bullet(doc, [n("3. Spatial-proximity merge ("), code_inline("_tinyCornersOfOneBill"),
                     n(") — catches the case where two corner-number boxes of the same bill survive NMS because their IoU is zero. Boxes covering < 25 % of the frame with the same canonical denomination get collapsed.")], level=1)
    add_bullet(doc, [n("4. Whitelist filter — anything outside the loaded labels file is dropped.")], level=1)
    add_bullet(doc, [n("5. 3-frame rolling vote ("), code_inline("voteFrames()"),
                     n(") — a denomination must appear in ≥ 2 of the last 3 frames to count, killing single-frame phantoms.")], level=1)
    add_bullet(doc, [n("The breakdown UI shows \"100 جنيه\" for a single bill, \"100 جنيه ×2\" for multiples; total is announced via TTS.")])

    add_h2(doc, "Object mode")
    add_para(doc, [n("What's in front of me?")])
    add_bullet(doc, [n("Multi-model ensemble in "), code_inline("multi_object_detector.dart"),
                     n(" running three YOLOs in series:")])
    add_bullet(doc, [code_inline("yolo_oiv7.tflite"), n(" — Open Images V7, 601 classes, specificity 30")], level=1)
    add_bullet(doc, [code_inline("yolo_ewaste.tflite"), n(" — 77 classes of electronics/e-waste, specificity 25")], level=1)
    add_bullet(doc, [code_inline("yolo_coco.tflite"), n(" — COCO, 80 classes, specificity 10")], level=1)
    add_bullet(doc, [n("Total "), b("758 effective classes"), n(" in the ensemble.")])
    add_bullet(doc, [n("Cross-model NMS: detections sorted by (specificity desc, confidence desc), then any later detection with IoU ≥ 0.45 vs. a kept one is absorbed. Top-K = 20.")])
    add_bullet(doc, [n("Falls back to single-bundle YoloObjectDetector if the multi bundle fails to load; falls back further to ML Kit's "),
                     code_inline("google_mlkit_object_detection"), n(" if all TFLite models fail.")])
    add_bullet(doc, [n("Each detected object speaks its label via TTS on first detection.")])

    add_h2(doc, "Scene / labels mode")
    add_para(doc, [n("Ambient description via "), code_inline("google_mlkit_image_labeling"),
                   n(". Useful for context that doesn't have a single dominant object.")])

    add_h2(doc, "Text mode")
    add_para(doc, [n("OCR via "), code_inline("google_mlkit_text_recognition"),
                   n(". Picks up Arabic + Latin text from signs, menus, receipts.")])

    add_h2(doc, "Pose mode")
    add_para(doc, [code_inline("google_mlkit_pose_detection"),
                   n(". Mostly used inside the sign translator; exposed in Vision for accessibility experiments.")])

    add_para(doc, [
        n("All modes share the same input source (live camera via "), code_inline("camera"),
        n(", or still photo via "), code_inline("image_picker"),
        n("), the same TTS announcer, and the same connectivity-aware UI (offline indicator if "),
        code_inline("connectivity_plus"),
        n(" reports no link). Recent results cache to Hive ("),
        code_inline("hive_flutter"), n(") for offline replay."),
    ])

    # ── §4 Safety / SOS ─────────────────────────────────────────────────
    add_h1(doc, "4. Safety / SOS")
    add_para(doc, [n("Lives at "), code_inline("/safety"),
                   n(" with a fullscreen modal at "), code_inline("/sos"), n(".")])
    add_bullet(doc, [n("Big SOS button arms a "), b("5-second cancellable countdown"),
                     n(" ("), code_inline("HapticFeedback.heavyImpact"), n(" × 3 on press).")])
    add_bullet(doc, [n("During countdown an opaque overlay covers the screen with a giant Cancel button.")])
    add_bullet(doc, [n("On expiry: "), code_inline("SosCoordinator.dispatch()"),
                     n(" fans out the alert through every configured channel for every contact:")])
    add_bullet(doc, [b("SMS"), n(" via "), code_inline("another_telephony"),
                     n(" (Android only, requires permission).")], level=1)
    add_bullet(doc, [b("WhatsApp"), n(" message via Twilio.")], level=1)
    add_bullet(doc, [b("Telegram"), n(" via bot.")], level=1)
    add_bullet(doc, [n("Channels report success/failure per-contact in real time; the result is also spoken via TTS.")])
    add_bullet(doc, [n("If geolocation is permitted ("), code_inline("geolocator"),
                     n("), the message embeds a maps link with current lat/lon and reverse-geocoded address.")])
    add_bullet(doc, [n("If no contacts exist, the button routes to "), code_inline("/profile/contacts"),
                     n(" with a snackbar prompt rather than firing into the void.")])
    add_bullet(doc, [b("Hardware trigger"), n(": an ESP32 SOS device paired through "),
                     code_inline("/hardware-pairing"), n(" keeps a WebSocket open ("),
                     code_inline("web_socket_channel"),
                     n(") and can fire the same dispatch path while the app is backgrounded.")])
    add_bullet(doc, [n("All dispatches log to a server-side history that the user can review under their account.")])
    add_para(doc, [b("Emergency contacts"), n(" at "), code_inline("/profile/contacts"),
                   n(" — full CRUD against the server (/api/contacts), per-channel enable flags (SMS / WhatsApp / Telegram), drag-to-reorder priority.")])

    # ── §5 Learning ─────────────────────────────────────────────────────
    add_h1(doc, "5. Learning hub")
    add_para(doc, [n("Lives at "), code_inline("/learning"), n(".")])
    add_bullet(doc, [n("Lesson list at "), code_inline("/learning/lesson/:id"),
                     n(" — each lesson has bundled video, ESL-clip examples, written explanation, and a Quiz call-to-action.")])
    add_bullet(doc, [b("Lesson quiz"), n(" at "), code_inline("/learning/lesson-quiz/:id"),
                     n(" — multiple-choice scoped to one lesson's vocabulary.")])
    add_bullet(doc, [b("Free quiz"), n(" at "), code_inline("/learning/quiz"),
                     n(" — cross-lesson randomized drill. Progress and best scores persist to the server so streaks survive reinstalls.")])
    add_bullet(doc, [n("Lesson video player uses "), code_inline("video_player"),
                     n(" + "), code_inline("chewie"),
                     n(" for chrome (fullscreen, scrubber, speed control). YouTube embeds use "),
                     code_inline("youtube_player_flutter"), n(".")])
    add_bullet(doc, [n("Reading-progress and visibility tracking via "), code_inline("visibility_detector"),
                     n(" — lessons are marked complete when ≥ 80 % scrolled.")])

    # ── §6 Shop ─────────────────────────────────────────────────────────
    add_h1(doc, "6. Shop")
    add_para(doc, [n("Lives at "), code_inline("/shop"),
                   n(", product detail at "), code_inline("/shop/product/:id"),
                   n(", cart at "), code_inline("/shop/cart"), n(".")])
    add_bullet(doc, [n("Product catalog backed by "), code_inline("assets/products/"),
                     n(" (local JSON + images) and overlaid by server-fetched live inventory.")])
    add_bullet(doc, [n("Skeleton loaders via "), code_inline("shimmer"),
                     n(" while data lands; product images via "), code_inline("cached_network_image"),
                     n(" for offline reload.")])
    add_bullet(doc, [n("Cart is local-first (Hive) and reconciles with the server on checkout open.")])
    add_bullet(doc, [n("Charts on the shop dashboard via "), code_inline("fl_chart"), n(".")])

    # ── §7 Assistant ────────────────────────────────────────────────────
    add_h1(doc, "7. Assistant (chatbot)")
    add_para(doc, [n("Lives at "), code_inline("/assistant"), n(".")])
    add_bullet(doc, [n("Conversational AI tab for plain Q&A — wired to a Gemini-backed server endpoint.")])
    add_bullet(doc, [n("Suggested-prompt chips matching the user's current context (Vision → \"What can the camera recognize?\", Safety → \"Walk me through SOS\").")])
    add_bullet(doc, [n("Supports both Arabic and English.")])
    add_bullet(doc, [n("Falls back to a cached \"FAQ\" knowledge base if the network call fails.")])

    # ── §8 Profile & Settings ───────────────────────────────────────────
    add_h1(doc, "8. Profile & settings")
    add_para(doc, [n("Lives at "), code_inline("/profile"), n(" with sub-routes.")])

    add_h2(doc, "Accessibility settings — every toggle has a measurable effect")
    for line in [
        "Auto-TTS on screen mount.",
        "Text scale slider (wired to MediaQuery.textScaler).",
        "High-contrast palette (WCAG AAA borders, swapped semantic colours).",
        "Color-blind palettes (deuter / protan / tritan-safe pairs across success / warning / error).",
        "Dyslexia font (bundled OpenDyslexic, swaps textTheme.fontFamily).",
        "Motor mode (enforces 56×56 min hit-target via MotorAware widget; bumps VisualDensity).",
        "Reduce motion (collapses flutter_animate transitions to instant, disables hero animations).",
        "Haptics-on-every-action (central HapticsService interceptor).",
        "Prefer sign language (translator opens in sign→text by default, Learning Hub auto-plays the first matching clip).",
        "A live preview screen renders each toggle's visible effect side-by-side.",
    ]:
        add_bullet(doc, [n(line)])

    add_h2(doc, "Follow Ishara")
    add_para(doc, [n("At "), code_inline("/profile/social"),
                   n(" — read-only screen with brand handles (Instagram, Facebook, Twitter, TikTok, YouTube, WhatsApp Business). Tapping opens the native app via "),
                   code_inline("url_launcher"),
                   n(". Sourced from "), code_inline("lib/src/core/config/ishara_brand.dart"), n(".")])

    add_h2(doc, "Emergency contacts")
    add_para(doc, [n("At "), code_inline("/profile/contacts"), n(" (see Section 4).")])

    add_h2(doc, "Language toggle (Arabic / English)")
    add_para(doc, [n("Flips "), code_inline("translations.dart"),
                   n(" AND wraps the app in a "), code_inline("Directionality"),
                   n(" whose textDirection is locale-derived, so Arabic mode is RTL across every screen.")])

    # ── §9 Hardware pairing ─────────────────────────────────────────────
    add_h1(doc, "9. Hardware pairing (ESP32)")
    add_para(doc, [n("Lives at "), code_inline("/hardware-pairing"), n(".")])
    add_bullet(doc, [n("WebSocket discovery ("), code_inline("web_socket_channel"),
                     n(") and pairing flow for an ESP32 glasses prototype (camera + IMU + emergency button).")])
    add_bullet(doc, [n("Streams frames over Wi-Fi back to Vision; streams accelerometer events to the safety service (fall detection → SOS arm).")])
    add_bullet(doc, [n("Device info exposed via "), code_inline("device_info_plus"),
                     n("; vibration confirmation on pair via "), code_inline("vibration"), n(".")])

    # ── §10 Cross-cutting ───────────────────────────────────────────────
    add_h1(doc, "10. Cross-cutting infrastructure")
    add_bullet(doc, [b("State management"), n(": Riverpod for everything, plus a few legacy provider consumers. GoRouter for navigation with a single Auth-aware refresh listenable.")])
    add_bullet(doc, [b("Theming"), n(": Custom IsharaColors token system (teal/orange dual-accent, light/dark, glassmorphism cards, pill radii, min touch targets). Token values are driven by accessibility settings so colour-blind palettes propagate everywhere.")])
    add_bullet(doc, [b("i18n"), n(": "), code_inline("translations.dart"),
                     n(" plus "), code_inline("intl"),
                     n(" formatters. Arabic is the default locale; English is fully supported. Directionality flips on locale change without app restart.")])
    add_bullet(doc, [b("Offline cache"), n(": Hive for product catalog, recent vision results, lesson progress, contacts; "),
                     code_inline("cached_network_image"),
                     n(" for product images; PSL sign clips stream from GitHub Releases (cached by the OS HTTP cache).")])
    add_bullet(doc, [b("Notifications"), n(": "), code_inline("flutter_local_notifications"),
                     n(" for SOS arming countdown reminders and lesson reminders.")])
    add_bullet(doc, [b("Animations"), n(": "), code_inline("flutter_animate"),
                     n(" for entrance/slide effects, "), code_inline("lottie"),
                     n(" for richer empty states.")])
    add_bullet(doc, [b("Share & device"), n(": "), code_inline("share_plus"),
                     n(" (translator results, vision OCR panel, learning word sheet, SOS history), "),
                     code_inline("permission_handler"), n(", "), code_inline("device_info_plus"),
                     n(", "), code_inline("connectivity_plus"), n(".")])
    add_bullet(doc, [b("Web compatibility"), n(": TFLite is FFI-only on mobile; the conditional import at "),
                     code_inline("lib/src/features/translator/data/tflite_flutter_web_stub.dart"),
                     n(" provides no-op stubs so the web build compiles. All non-camera features (translator text mode, learning, shop, assistant) work on web; camera-bound features show \"Mobile only\" tiles.")])

    # ── §11 Model inventory ─────────────────────────────────────────────
    add_h1(doc, "11. On-device model inventory")
    add_table(doc,
        ["File", "Classes", "Size", "Purpose"],
        [
            ["sign_yolo.tflite",     "10",  "~6.2 MB",  "Arabic Sign Language YOLOv8n (mAP50 0.995)"],
            ["currency_egp.tflite",  "6",   "~6.2 MB",  "EGP banknote YOLOv8n (mAP50 0.995)"],
            ["yolo_coco.tflite",     "80",  "~6.5 MB",  "COCO objects YOLOv8n"],
            ["yolo_oiv7.tflite",     "601", "~7.2 MB",  "Open Images V7 fine-grained YOLOv8n"],
            ["yolo_ewaste.tflite",   "77",  "~12.1 MB", "Electronics / e-waste YOLOv8n (mAP50 0.33)"],
        ],
        col_widths=[1.8, 0.8, 1.0, 2.8],
    )
    add_para(doc, [n("Plus matching _labels.txt files for each model.")])

    # ── Save ────────────────────────────────────────────────────────────
    doc.save(OUT)
    print(f"[done] {OUT} ({OUT.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    build()
